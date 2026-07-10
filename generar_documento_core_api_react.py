from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUT = Path("Documento_Especificacion_Actualizado_Core_API_React.docx")


doc = Document()
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(10)


def h(text, level=1):
    doc.add_heading(text, level=level)


def p(text=""):
    doc.add_paragraph(text)


def bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def code(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover.add_run("Documento de Especificacion Funcional y Tecnica Actualizado\n")
run.bold = True
run.font.size = Pt(16)
cover.add_run("Sistema de Analisis de Degradacion de Equipos TI\n")
cover.add_run("Core API + React\n")
cover.add_run("Emil Encalada - Jorge Ramos\n")
cover.add_run("Actualizado: 09/07/2026")
doc.add_page_break()

h("1. Descripcion general")
p(
    "El sistema integra un backend ASP.NET Core que expone la API del Core de "
    "degradacion y un frontend React con Vite que consume dicha API. El archivo "
    "ZIP recibido corresponde al frontend React, no a la API. La API se encuentra "
    "implementada dentro del proyecto .NET Proyecto_EmilEncalada.Server."
)
p(
    "La solucion permite administrar departamentos, equipos y reparaciones, y "
    "evaluar la degradacion de equipos tecnologicos a partir de su antiguedad, "
    "vida util, tipo, estado, historial de reparaciones y costos asociados."
)

h("2. Componentes del sistema")
bullets(
    [
        "Backend/API: Proyecto_EmilEncalada.Server, desarrollado en ASP.NET Core.",
        "Frontend React: core_degradacion_react, desarrollado con React, Vite, Axios, Bootstrap y React Router.",
        "Base de datos: SQL Server online configurado desde appsettings.json.",
        "Core: logica de calculo de degradacion usando Strategy, Repository, Service Layer y DTOs.",
    ]
)

h("3. Donde esta implementada la API del Core")
p("La API del Core esta implementada en el backend .NET, especialmente en el controlador CoreController:")
code("Proyecto_EmilEncalada.Server/Controllers/CoreController.cs")
p("Endpoint principal:")
code("GET /api/Core/degradacion/{id}")
p("URL local de ejecucion:")
code("http://localhost:5086/api/Core/degradacion/1")
p("Swagger local:")
code("http://localhost:5086/swagger")

h("4. Flujo del Core")
bullets(
    [
        "React o Swagger llama a GET /api/Core/degradacion/{id}.",
        "CoreController recibe la peticion HTTP.",
        "CoreController delega el calculo a DegradacionService.",
        "DegradacionService obtiene el equipo y sus reparaciones mediante IEquipoRepository.",
        "EquipoRepository consulta la base de datos mediante Entity Framework.",
        "DegradacionService entrega los datos a ICalculoDegradacionStrategy.",
        "CalculoDegradacionPorReglasStrategy calcula porcentaje de degradacion, nivel, valor actual y recomendacion.",
        "El backend devuelve un ResultadoDegradacionDto en formato JSON.",
    ]
)

h("5. Patrones de diseno y SOLID")
h("5.1 Strategy", 2)
p(
    "El patron Strategy permite encapsular la formula de calculo de degradacion en "
    "una clase independiente. Esto permite agregar nuevas estrategias sin modificar "
    "el controlador ni el servicio."
)
code("Interfaces/ICalculoDegradacionStrategy.cs\nStrategies/CalculoDegradacionPorReglasStrategy.cs")

h("5.2 Repository", 2)
p(
    "El patron Repository encapsula el acceso a datos de equipos y reparaciones "
    "para que el Core no dependa directamente de Entity Framework."
)
code("Interfaces/IEquipoRepository.cs\nRepositories/EquipoRepository.cs")

h("5.3 Service Layer", 2)
p("DegradacionService coordina el repositorio y la estrategia de calculo.")
code("Services/DegradacionService.cs")

h("5.4 DTO", 2)
p("ResultadoDegradacionDto define el contrato JSON devuelto por la API del Core.")
code("Dtos/ResultadoDegradacionDto.cs")

h("5.5 SOLID aplicado", 2)
bullets(
    [
        "SRP: CoreController, DegradacionService, EquipoRepository y la Strategy tienen responsabilidades separadas.",
        "OCP: nuevas formulas pueden agregarse creando nuevas clases que implementen ICalculoDegradacionStrategy.",
        "DIP: DegradacionService depende de interfaces y no de implementaciones concretas.",
    ]
)

h("6. API JSON implementada")
p("Endpoints administrativos y del Core:")
code(
    "GET    /api/Departamentos\n"
    "POST   /api/Departamentos\n"
    "GET    /api/Equipos\n"
    "POST   /api/Equipos\n"
    "GET    /api/Reparaciones\n"
    "POST   /api/Reparaciones\n"
    "GET    /api/Core/degradacion/{id}"
)
p("Para crear, editar o eliminar datos, el frontend envia el header:")
code("X-Rol: Admin")

h("7. Ejemplo de respuesta del Core")
code(
    '{\n'
    '  "equipo": "Laptop Oficina",\n'
    '  "tipo": "Laptop",\n'
    '  "marca": "Dell",\n'
    '  "estado": "Regular",\n'
    '  "antiguedadMeses": 24.02,\n'
    '  "vidaUtilMeses": 60,\n'
    '  "vidaRestanteMeses": 35.98,\n'
    '  "costoInicial": 1000,\n'
    '  "cantidadReparaciones": 2,\n'
    '  "costoReparaciones": 250,\n'
    '  "porcentajeCostoReparacion": 25.00,\n'
    '  "degradacionPorcentaje": 58.35,\n'
    '  "nivelDegradacion": "Media",\n'
    '  "valorActualEstimado": 416.50,\n'
    '  "recomendacion": "Se recomienda dar seguimiento y planificar mantenimiento."\n'
    '}'
)

h("8. Frontend React integrado")
p("El frontend React integrado desde el ZIP se ubica en:")
code("core_degradacion_react")
p("Consume el backend local mediante Axios en:")
code("core_degradacion_react/src/api/axiosConfig.js")
p("URL local del API usada por React:")
code("VITE_API_URL=http://localhost:5086")
p("Rutas React:")
bullets(
    [
        "/: pantalla de inicio.",
        "/degradacion: calculo individual de degradacion por ID.",
        "/analisis: analisis avanzado con matriz de escenarios.",
        "/gestion: gestion de datos para agregar, listar y evaluar departamentos, equipos y reparaciones.",
    ]
)

h("9. URLs locales de ejecucion")
code(
    "Backend/API: http://localhost:5086\n"
    "Swagger: http://localhost:5086/swagger\n"
    "React: http://localhost:5174\n"
    "Gestion de datos React: http://localhost:5174/gestion"
)

h("10. Base de datos")
p(
    "El backend usa SQL Server online mediante la cadena de conexion configurada "
    "en appsettings.json. Adicionalmente, se implemento DatabaseBootstrapper para "
    "crear o completar automaticamente tablas necesarias al iniciar el backend."
)
code("Proyecto_EmilEncalada.Server/Data/DatabaseBootstrapper.cs")
p("Tablas principales:")
bullets(["Usuarios", "Departamentos", "Equipos", "Reparaciones", "Productos"])

h("11. Despliegue")
p("Para que el sistema funcione en produccion se deben desplegar dos piezas:")
bullets(
    [
        "Backend/API ASP.NET Core en un hosting compatible con .NET.",
        "Frontend React como sitio estatico generado con Vite.",
    ]
)
p("En despliegue, React no debe usar localhost. Debe apuntar a la URL publica del backend:")
code("VITE_API_URL=https://URL-PUBLICA-DE-TU-API")
p(
    "Si se despliega solo React sin desplegar la API, la interfaz cargara, pero "
    "no podra listar, agregar ni evaluar datos."
)

h("12. Comandos locales")
p("Backend:")
code("dotnet run --project Proyecto_EmilEncalada.Server\\Proyecto_EmilEncalada.Server.csproj --launch-profile https")
p("React:")
code("cd core_degradacion_react\nnpm run dev -- --host 0.0.0.0")
p("Build React:")
code("cd core_degradacion_react\nnpm run build")

h("13. Conclusion")
p(
    "La solucion cumple con la separacion entre API/Core y frontend React. La API "
    "se encuentra implementada en ASP.NET Core y aplica Strategy, Repository, "
    "Service Layer, DTOs y principios SOLID. React consume esa API mediante Axios "
    "y permite ejecutar el calculo individual, analisis avanzado y gestion de datos. "
    "Para despliegue real, el backend debe publicarse con una URL accesible y React "
    "debe configurarse para consumir esa URL publica."
)

doc.save(OUT)
print(OUT.resolve())
